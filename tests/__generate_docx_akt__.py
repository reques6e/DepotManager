import json
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches

def set_col_width(column, width):
    """
    Функция для установки ширины столбца.
    width задается в точках (1 pt = 1/72 дюйма).
    """
    for cell in column.cells:
        cell.width = width

def add_quantity_suffix(quantity):
    """
    Функция для добавления "шт" к количеству, если его нет в строке.
    """
    # Преобразуем количество в строку
    quantity_str = str(quantity)
    
    if "шт" not in quantity_str:
        quantity_str += " шт"
    
    return quantity_str

def generate_document(data):
    # Создание документа
    doc = Document()

    # Установка шрифта по умолчанию
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(12)

    # Добавление заголовка
    header = doc.add_paragraph()
    header.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    header.add_run(data['header']['company']).bold = True

    # Добавление barcode_data справа, серым цветом
    barcode_data = data.get("barcode_data", "Error")  # Если нет данных о штрих-коде, используем заглушку
    run = header.add_run("\t\t\t" + 'ID: ' + barcode_data)
    run.font.color.rgb = RGBColor(128, 128, 128)  # Серый цвет

    # Добавление основного заголовка
    main_header = doc.add_paragraph()
    main_header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    main_header.add_run(data['header']['act']).bold = True

    # Добавление таблицы с организационной информацией
    org_table = doc.add_table(rows=3, cols=2)
    org_table.style = None  # Убираем стиль с границами

    # Установим ширину ячеек
    set_col_width(org_table.columns[0], 200)  # Левая колонка шире
    set_col_width(org_table.columns[1], 100000000)  # Правая колонка

    # Заполнение данных таблицы
    for i, (left_text, right_text) in enumerate(data['org_data']):
        # Левая колонка серым текстом
        cell = org_table.cell(i, 0)
        cell.text = left_text
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run = paragraph.runs[0]
            run.font.color.rgb = RGBColor(128, 128, 128)  # Серый цвет
            run.font.size = Pt(10)

        # Правая колонка обычным текстом
        org_table.cell(i, 1).text = right_text
        for paragraph in org_table.cell(i, 1).paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            for run in paragraph.runs:
                run.font.size = Pt(10)

    # Отступ перед следующей таблицей
    doc.add_paragraph()

    # Таблица товаров
    product_table = doc.add_table(rows=len(data['product_data']) + 1, cols=4)
    product_table.style = 'Table Grid'

    # Установим ширину ячеек для "№", "Артикул" и "Кол-во"
    set_col_width(product_table.columns[0], 20)   # Узкое поле для "№"
    set_col_width(product_table.columns[1], 1090000)   # Увеличиваем ширину для "Артикул"
    set_col_width(product_table.columns[2], 9900000)   # Увеличиваем ширину для "Товар"
    set_col_width(product_table.columns[3], 900000)   # Узкое поле для "Кол-во"

    # Заголовок таблицы
    for i, header_text in enumerate(data['product_data_headers']):
        cell = product_table.cell(0, i)
        cell.text = header_text
        for paragraph in cell.paragraphs:
            if header_text == "Товар":
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # По центру для "Товар"
            else:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # По левому краю для остальных
            run = paragraph.runs[0]
            run.font.size = Pt(10)
            run.font.bold = True

    # Данные в таблице
    for row_idx, row_data in enumerate(data['product_data']):
        for col_idx, text in enumerate(row_data):
            # Если это столбец "Кол-во", добавим "шт", если его нет
            if col_idx == 3:
                text = add_quantity_suffix(text)
            cell = product_table.cell(row_idx + 1, col_idx)
            cell.text = text
            for paragraph in cell.paragraphs:
                if col_idx == 3:  # "Кол-во"
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # По центру для "Кол-во"
                else:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # По левому краю для остальных
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    # Добавление подписей
    doc.add_paragraph()
    signatures = doc.add_paragraph()
    for signature in data['signatures']:
        signatures.add_run(signature + '\n')

    # Сохранение документа
    doc.save('Акт_о_списании.docx')

if __name__ == "__main__":
    with open('data.json', 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    generate_document(data)
